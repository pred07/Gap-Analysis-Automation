#!/usr/bin/env python3
"""
Document Parser for Batch Analysis

Parses Excel, PDF, and DOCX files to extract security-relevant information
including URLs, endpoints, control checklists, and policy documentation.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Dict, List, Optional, Any
import logging

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


logger = logging.getLogger(__name__)


class DocumentParser:
    """Parse security documents (Excel, PDF, DOCX) for batch analysis."""
    
    def __init__(self, debug: bool = False):
        """
        Initialize document parser.
        
        Args:
            debug: Enable debug logging
        """
        self.debug = debug
        if debug:
            logger.setLevel(logging.DEBUG)
        
        # URL pattern for extraction
        self.url_pattern = re.compile(
            r'https?://[^\s<>"{}|\\^`\[\]]+',
            re.IGNORECASE
        )
        
        # Control ID pattern (e.g., CTRL-001, Control_1, etc.)
        self.control_pattern = re.compile(
            r'(?:CTRL|Control|CTL)[-_\s]?\d+',
            re.IGNORECASE
        )
    
    def parse_directory(self, directory: Path) -> Dict[str, Any]:
        """
        Parse all supported documents in a directory.
        
        Args:
            directory: Path to directory containing documents
            
        Returns:
            Dictionary with parsed data from all documents
        """
        if not directory.exists():
            logger.warning(f"Directory not found: {directory}")
            return {"urls": [], "controls": [], "policies": [], "errors": []}
        
        results = {
            "urls": [],
            "controls": [],
            "policies": [],
            "metadata": [],
            "errors": []
        }
        
        # Supported file extensions
        extensions = {
            ".xlsx": self.parse_excel,
            ".xls": self.parse_excel,
            ".pdf": self.parse_pdf,
            ".docx": self.parse_docx,
        }
        
        for file_path in directory.rglob("*"):
            if not file_path.is_file():
                continue
            
            ext = file_path.suffix.lower()
            if ext in extensions:
                try:
                    logger.info(f"Parsing {file_path.name}...")
                    data = extensions[ext](file_path)
                    
                    # Merge results
                    results["urls"].extend(data.get("urls", []))
                    results["controls"].extend(data.get("controls", []))
                    results["policies"].extend(data.get("policies", []))
                    results["metadata"].append({
                        "file": file_path.name,
                        "type": ext,
                        "status": "success"
                    })
                    
                except Exception as e:
                    error_msg = f"Error parsing {file_path.name}: {e}"
                    logger.error(error_msg)
                    results["errors"].append(error_msg)
                    results["metadata"].append({
                        "file": file_path.name,
                        "type": ext,
                        "status": "failed",
                        "error": str(e)
                    })
        
        # Deduplicate URLs
        results["urls"] = list(set(results["urls"]))
        results["controls"] = list(set(results["controls"]))
        
        logger.info(f"Parsed {len(results['metadata'])} documents")
        logger.info(f"Found {len(results['urls'])} unique URLs")
        logger.info(f"Found {len(results['controls'])} unique controls")
        
        return results
    
    def parse_excel(self, file_path: Path) -> Dict[str, Any]:
        """
        Parse Excel file for security data.
        
        Args:
            file_path: Path to Excel file
            
        Returns:
            Dictionary with extracted data
        """
        if not EXCEL_AVAILABLE:
            raise ImportError("openpyxl not installed. Run: pip install openpyxl")
        
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        
        data = {
            "urls": [],
            "controls": [],
            "policies": []
        }
        
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows(values_only=True):
                for cell in row:
                    if cell is None:
                        continue
                    
                    cell_str = str(cell)
                    
                    # Extract URLs
                    urls = self.url_pattern.findall(cell_str)
                    data["urls"].extend(urls)
                    
                    # Extract control IDs
                    controls = self.control_pattern.findall(cell_str)
                    data["controls"].extend(controls)
                    
                    # Look for policy keywords
                    if any(keyword in cell_str.lower() for keyword in 
                           ["policy", "procedure", "guideline", "standard"]):
                        data["policies"].append(cell_str[:200])  # First 200 chars
        
        return data
    
    def parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        """
        Parse PDF file for security data.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Dictionary with extracted data
        """
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")
        
        reader = PdfReader(file_path)
        
        data = {
            "urls": [],
            "controls": [],
            "policies": []
        }
        
        full_text = ""
        for page in reader.pages:
            text = page.extract_text()
            full_text += text + "\n"
        
        # Extract URLs
        urls = self.url_pattern.findall(full_text)
        data["urls"].extend(urls)
        
        # Extract control IDs
        controls = self.control_pattern.findall(full_text)
        data["controls"].extend(controls)
        
        # Extract policy sections (paragraphs with policy keywords)
        lines = full_text.split("\n")
        for i, line in enumerate(lines):
            if any(keyword in line.lower() for keyword in 
                   ["policy", "procedure", "guideline", "standard"]):
                # Get context (current line + next 2 lines)
                context = " ".join(lines[i:i+3])
                data["policies"].append(context[:300])
        
        return data
    
    def parse_docx(self, file_path: Path) -> Dict[str, Any]:
        """
        Parse DOCX file for security data.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary with extracted data
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        doc = Document(file_path)
        
        data = {
            "urls": [],
            "controls": [],
            "policies": []
        }
        
        # Parse paragraphs
        for para in doc.paragraphs:
            text = para.text
            
            # Extract URLs
            urls = self.url_pattern.findall(text)
            data["urls"].extend(urls)
            
            # Extract control IDs
            controls = self.control_pattern.findall(text)
            data["controls"].extend(controls)
            
            # Look for policy content
            if any(keyword in text.lower() for keyword in 
                   ["policy", "procedure", "guideline", "standard"]):
                data["policies"].append(text[:300])
        
        # Parse tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    
                    # Extract URLs from tables
                    urls = self.url_pattern.findall(text)
                    data["urls"].extend(urls)
                    
                    # Extract controls from tables
                    controls = self.control_pattern.findall(text)
                    data["controls"].extend(controls)
        
        return data
    
    def extract_urls(self, content: Dict[str, Any]) -> List[str]:
        """
        Extract and validate URLs from parsed content.
        
        Args:
            content: Parsed document content
            
        Returns:
            List of valid URLs
        """
        urls = content.get("urls", [])
        
        # Filter and validate
        valid_urls = []
        for url in urls:
            # Basic validation
            if url.startswith(("http://", "https://")) and len(url) > 10:
                valid_urls.append(url)
        
        return list(set(valid_urls))
    
    def extract_controls(self, content: Dict[str, Any]) -> List[str]:
        """
        Extract control IDs from parsed content.
        
        Args:
            content: Parsed document content
            
        Returns:
            List of control IDs
        """
        return list(set(content.get("controls", [])))


if __name__ == "__main__":
    # Test the parser
    import sys
    
    logging.basicConfig(level=logging.INFO)
    
    if len(sys.argv) > 1:
        test_path = Path(sys.argv[1])
    else:
        test_path = Path("batch_inputs/documents")
    
    parser = DocumentParser(debug=True)
    results = parser.parse_directory(test_path)
    
    print("\n=== Parsing Results ===")
    print(f"URLs found: {len(results['urls'])}")
    print(f"Controls found: {len(results['controls'])}")
    print(f"Policies found: {len(results['policies'])}")
    print(f"Errors: {len(results['errors'])}")
    
    if results['urls']:
        print("\nSample URLs:")
        for url in results['urls'][:5]:
            print(f"  - {url}")
    
    if results['errors']:
        print("\nErrors:")
        for error in results['errors']:
            print(f"  - {error}")
